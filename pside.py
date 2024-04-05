from io import BytesIO
import streamlit as st
import pandas as pd
from pymongo import MongoClient
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from uuid import uuid4
from docx import Document
from openai import OpenAI
import os
from datetime import datetime
import re  # Import the regular expressions module
from st_copy_to_clipboard import st_copy_to_clipboard
import json


new_prompt = """\nPHASE 3- Documentation Synthesis:
Document the patient's responses and any additional relevant information, as a nurse would for a doctor to use.

*rules*
-You have examples of what a nurses notes look like. Follow the 'SOAP' model for writing notes, you also have access to this in your knowledge.
-You should be very careful filling the O-Objective part of the note as most times you will not have the information to fill this out.
-Keep the notes easy and quick to read for a doctor. Ensure the notes get created in a Microsoft word document.
-You MUST use your knowledge of how to write good nursing notes and nursing note examples to format your notes properly.
-You MUST create a word document of the notes to deliver your final notes.
- You MUST use BULLET POINTS to allow for quick reading by doctor. They should not be repetitive
- For the implementation tab, you MUST not talk about communicating to a healthcare provider because they are visiting a doctor already

Here is an example document:

Subjective: 
- Patient reports a dull ache in the left shoulder, which started about a month ago
- The pain intensifies during workouts, especially during bench press. No previous injuries reported.
- No swelling, redness, or warmth in the area. Patient mentions a clicking sound when raising the arm laterally.
Objective: 
- Patient has been using ice and ibuprofen for pain management. 
- Ibuprofen provides temporary relief, and is needed approximately twice a day.
Analysis: 
- The symptoms suggest a potential joint issue in the left shoulder, possibly related to exercise. 
- The clicking sound during movement may indicate a ligament or tendon problem.
Plan: 
- Possible imaging tests or referral to physical therapy may be considered based on the assessment.
Implementation:
- Advised the patient on the importance of consulting a healthcare provider
- Discussed the potential need for further diagnostic tests or physical therapy.
Evaluation: 
- To be determined based on patient's follow-up with healthcare provider and any subsequent treatment.

Make sure to include the implementation and evaluation.

Here is the chat history to base this off of below: \n"""

api_key = os.getenv('OPENAI_API_KEY')
openai_client = OpenAI(api_key=api_key)

# MongoDB setup
mongo_key = os.getenv('MONGO_KEY')
client = MongoClient('mongodb+srv://ishansheth31:Kevi5han1234@breezytest1.saw2kxe.mongodb.net/')


db = client.breezyaccounts
accounts_collection = db.accounts


import streamlit as st

def login_form():
    with st.form("login_form"):
        st.title("Physician Login")
        email = st.text_input("Email", key="login_email")
        password = st.text_input("Password", type="password", key="login_password")
        submit_button = st.form_submit_button("Login")
        return submit_button, email, password

def validate_credentials(email, password):
    user_doc = accounts_collection.find_one({"Email": email, "Password": password})
    if user_doc:
        # Retrieve the patients collection for the logged-in user
        main_db = user_doc["Database"]
        db1 = client[main_db]
        db_test = user_doc["Collection"]
        patients_collection = db1[db_test]
        st.session_state['logged_in'] = True
        st.session_state['patients_collection'] = patients_collection
        st.session_state['user_email'] = email  # Store the user's email in session state for easy access
        return patients_collection
    else:
        st.error("Invalid email or password. Please try again.")
        return None



def display_main_content(email, password):
    if validate_credentials(email, password):
        display_patient_info()  # Adjust if this function also needs the collection passed as an argument
    else:
        st.error("Failed to access patient data.")


def send_email(to_email, link, patients_collection, fname):
    patient_id = str(uuid4())
    # Modify the link to include the patient ID as a query parameter
    personalized_link = f"{link}?patient_id={patient_id}"
    
    smtp_server = "smtp.gmail.com"
    smtp_port = 587
    smtp_username = 'mainbreezy11@gmail.com'
    smtp_password = 'qgzp kaay irpm hqyq'
    from_email = smtp_username
    subject = "Your Virtual Nurse Assessment"

    # HTML email body
    body = f"""
    <html>
        <body>
            <p>Hello {fname},</p>
            <p>Welcome to Breezy. You will be conducting your patient assessment so we can understand your reason for visiting in advance.</p>
            <p>Please complete your <a href="{personalized_link}">assessment</a> before your appointment.</p>
            <p>We look forward to seeing you soon!</p>
        </body>
    </html>
    """
    
    msg = MIMEMultipart("alternative")
    msg['From'] = from_email
    msg['To'] = to_email
    msg['Subject'] = subject
    
    # Attach the HTML version
    msg.attach(MIMEText(body, 'html'))

    try:
        server = smtplib.SMTP(smtp_server, smtp_port)
        server.starttls()
        server.login(smtp_username, smtp_password)
        server.send_message(msg)
        server.quit()
        # Update patient status in MongoDB
        patients_collection.update_one({"Email": to_email}, {"$set": {"PatientID": patient_id, "Status": "Sent"}})
        return True
    except Exception as e:
        st.error(f"Failed to send email to {to_email}: {e}")
        return False


def update_patient_status(patient_id, status, patients_collection):
    # Function to update the patient's status in MongoDB
    patients_collection.update_one({"PatientID": patient_id}, {"$set": {"Status": status}})

def check_and_update_patient_completion_status(patient_id, patients_collection):
    # Function to check if the assessment is completed and update status accordingly
    document = patients_collection.find_one({"PatientID": patient_id})
    if document and document.get("Assessment"):
        update_patient_status(patient_id, "Completed", patients_collection)
        return "Completed"
    return document.get("Status", "Not Sent")

def fetch_patients(patients_collection):
    patients_list = list(patients_collection.find())
    # Convert to DataFrame
    patients_df = pd.DataFrame(patients_list)
    
    # Check if 'Date' column exists
    if 'Date' in patients_df.columns:
        patients_df['Date'] = pd.to_datetime(patients_df['Date'])
    else:
        # Handle case where 'Date' does not exist
        # Option 1: Set a default date
        patients_df['Date'] = pd.to_datetime('1-31-2004')
        
        # Option 2: Drop rows without 'Date' or handle it according to your logic
        # For this example, we'll print a message and return an empty DataFrame
        return pd.DataFrame()  # Return an empty DataFrame or handle as needed
    
    return patients_df


def generate_patient_assessment_report(patient_id, patients_collection):
    document = patients_collection.find_one({"PatientID": patient_id})
    if not document:
        return "No document found with PatientID: " + patient_id

    assessment = document.get('Assessment', None)
    if not assessment:
        return "Assessment not found in document"

    formatted_chat_history = "\n".join([f'{entry[0]}: "{entry[1]}"' for entry in assessment])

    prompt = f"{new_prompt}\n\n{formatted_chat_history}"

    try:
        response = openai_client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "system", "content": prompt}]
        )
        generated_report = response.choices[0].message.content

        # Parse the generated report text to get the SOAP sections
        report_sections = parse_report_sections(generated_report)

        # Store the parsed sections in the database
        update_patient_report_in_db(patient_id, report_sections, patients_collection)

        return generated_report, report_sections
    except Exception as e:
        return f"An error occurred: {str(e)}"

def parse_report_sections(report_text):
    patterns = {
        "Subjective": r"Subjective:\s*(.*?)\n(?=Objective:)",
        "Objective": r"Objective:\s*(.*?)\n(?=Analysis:)",
        "Analysis": r"Analysis:\s*(.*?)\n(?=Plan:)",
        "Plan": r"Plan:\s*(.*?)\n(?=Implementation:)",
        "Implementation": r"Implementation:\s*(.*?)\n(?=Evaluation:)",
        "Evaluation": r"Evaluation:\s*(.*)"
    }
    sections = {}
    for key, pattern in patterns.items():
        match = re.search(pattern, report_text, re.DOTALL)
        if match:
            sections[key] = match.group(1).strip()
    return sections

def update_patient_report_in_db(patient_id, report_sections, patients_collection):
    if report_sections:
        patients_collection.update_one(
            {"PatientID": patient_id},
            {"$set": {section: text for section, text in report_sections.items()}}
        )


def add_new_patient(fName, lName, email, appointment_datetime, patients_collection):
    new_patient = {
        "fName": fName,
        "lName": lName,
        "Email": email,
        "Status": "Not Sent",
        "Date": appointment_datetime,
        "PatientID": "",  # Empty PatientID field
        "Assessment": ""
    }
    patients_collection.insert_one(new_patient)
    st.sidebar.success("Patient Added Successfully")
    st.experimental_rerun()  # Rerun the app to reflect the update immediately


def is_valid_email(email):
    # Simple regex pattern for validating an email
    pattern = r'^[a-z0-9._%+-]+@[a-z0-9.-]+\.[a-z]{2,}$'
    return re.match(pattern, email, re.I)  # re.I is for case-insensitive matching

def generate_downloadable_docx(patient_id, patients_collection):
    report_content, report_sections = generate_patient_assessment_report(patient_id, patients_collection)
    
    doc = Document()
    doc.add_heading('Patient Assessment Report', 0)
    
    for section, content in report_sections.items():
        doc.add_heading(section, level=1)
        doc.add_paragraph(content)
    
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return report_sections, doc_io


def display_patient_info():
    if not st.session_state.get('logged_in', False):
        submit_button, email, password = login_form()
        if submit_button:
            if validate_credentials(email, password) is not None:
                display_patient_data(st.session_state['patients_collection'])
                st.experimental_rerun()
            else:
                st.error("Login failed. Please check your credentials and try again.")
    else:
        # Directly display patient data if already logged in
        display_patient_data(st.session_state['patients_collection'])

        # Sidebar form for adding a new patient
        with st.sidebar.form("new_patient_form", clear_on_submit=True):  # clear_on_submit can clear the form after submission
            st.write("Add New Patient")
            fName = st.text_input("First Name")
            lName = st.text_input("Last Name")
            email = st.text_input("Email")
            appointment_date = st.date_input("Appointment Date")
            appointment_time = st.time_input("Appointment Time")
            submit_new_patient = st.form_submit_button("Add")

            if submit_new_patient:
                # Perform input validation here (e.g., check for empty fields)
                if fName and lName and email and appointment_date and appointment_time:
                    appointment_datetime = datetime.combine(appointment_date, appointment_time)
                    add_new_patient(fName, lName, email, appointment_datetime, st.session_state['patients_collection'])
                else:
                    st.sidebar.error("Please fill out all fields.")

        # Logout button at the bottom of the sidebar
        with st.sidebar:
            st.write("")  # Add some space
            if st.button("Logout"):
                # Clear session state
                for key in list(st.session_state.keys()):
                    del st.session_state[key]
                st.experimental_rerun()



def display_patient_data(patients_collection):
    st.title("Breezy Portal")
    st.subheader("**Today's Patients**")
    st.write("Be sure to press the 'R' key to refresh the page whenever a patient has completed their assessment. This way you can download it.")

    patients_df = fetch_patients(patients_collection)

    if not patients_df.empty:
        patients_df['Date'] = pd.to_datetime(patients_df['Date'], errors='coerce')
        patients_df.dropna(subset=['Date'], inplace=True)
        patients_df.sort_values(by='Date', inplace=True)

        min_date = patients_df['Date'].min().date()
        max_date = patients_df['Date'].max().date()
        selected_date = st.sidebar.date_input("Select a Date", min_value=min_date, max_value=max_date, value=min_date)

        selected_patients = patients_df[patients_df['Date'].dt.date == selected_date]

        for index, patient in selected_patients.iterrows():
            appointment_time = patient['Date'].strftime('%I:%M %p')
            with st.expander(f"{patient['fName']} {patient['lName']} - {appointment_time}"):
                st.write(f"Email: {patient['Email']}")
                patient_status = check_and_update_patient_completion_status(patient["PatientID"], patients_collection)
                st.write(f"Status: {patient_status}")

                email_sent_key = f"email_sent_{patient['PatientID']}"

                # Check if the email status is not 'Sent', and session state doesn't already indicate email sent
                if patient_status != "Sent" and not st.session_state.get(email_sent_key):
                    send_button = st.button("Send Email", key=patient["PatientID"])
                    if send_button:
                        link = "https://breezy.streamlit.app/"
                        email_sent = send_email(patient['Email'], link, patients_collection, patient['fName'])
                        if email_sent:
                            st.session_state[email_sent_key] = True  # Mark as sent in the session state
                            update_patient_status(patient["PatientID"], "Sent", patients_collection)  # Update status in the DB
                            st.success("Email sent successfully.")
                            st.experimental_rerun()
                        else:
                            st.error("Failed to send email.")

                if patient_status == "Completed":
                    # Check if the report already exists in the database
                    document = patients_collection.find_one({"PatientID": patient["PatientID"]})
                    if 'Subjective' in document and 'Objective' in document:
                        report_sections = {key: document[key] for key in ['Subjective', 'Objective', 'Analysis', 'Plan', 'Implementation', 'Evaluation']}
                        doc_io = generate_report_docx(report_sections)
                    else:
                        report_sections, doc_io = generate_downloadable_docx(patient["PatientID"], patients_collection)
                    
                    if report_sections:
                        for section, content in report_sections.items():
                            st.subheader(section)
                            st.markdown(content, unsafe_allow_html=True)
                            st_copy_to_clipboard(content)

                        full_name = f"{patient['fName']}_{patient['lName']}".replace(' ', '_')
                        file_name = f"{full_name}_Patient_Assessment_Report.docx"

                        st.download_button(label="Download as Word Document",
                                           data=doc_io.getvalue(),
                                           file_name=file_name,
                                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


def generate_report_docx(report_sections):
    doc = Document()
    doc.add_heading('Patient Assessment Report', 0)
    
    for section, content in report_sections.items():
        doc.add_heading(section, level=1)
        doc.add_paragraph(content)
    
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io

def main():
    display_patient_info()

if __name__ == "__main__":
    main()