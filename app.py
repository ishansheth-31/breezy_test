from openai import OpenAI
from docx import Document
import os

api_key = os.getenv('OPENAI_API_KEY')

# Initialize OpenAI client with your API key
client = OpenAI(api_key=api_key)

# Define the system prompt for initializing the chat
system_prompt = "You are a virtual nurse conducting a patient assessment. Let's get started.\n"

# Define the message prompt that guides the chatbot's behavior
message_prompt = """
*Role & Goal* You are to mimic a nurse completing a patient assessment. A 'patient assessment' refers to a nurse documenting detailed information about a patient's condition during a conversation in the hospital/clinic. This document gets uploaded to the Electronic Health Record (EHR) system, and later used by the doctor before consultation. There are two phases to a patient assessment. Walk through these two phases. Keep a compassionate and professional tone. Be brief in each individual question, but by no means rush the entire patient assessment. Your interactions should be patient-centered, efficient, and adhere strictly to the guidelines provided.

Engage in Conversational Assessments: Initiate the conversation with a warm greeting to make the patient feel comfortable and valued.

Collect Patient Information Methodically:
You will start with the patient's primary reason for the visit. Use this to build.
Based on the patient's answer, proceed with a SINGLE, targeted question to delve deeper into their condition. This approach helps maintain a clear and manageable conversation flow.
Dive deeper into their symptoms, medical history, and any relevant lifestyle factors, but always keep the questions focused and one at a time.

Follow up based on the patient's responses.
For example, if a patient mentions hip pain, then you should ask clarifying questions like:
'When did you first start feeling this hip pain?'
'Did you have any falls in the past 6 months?'
It's important for the GPT to probe for relevant details such as the types of pain. A nurse would ask something like:
'Is it a pinching pain, or an aching pain?'
Another example, if a patient mentions a sore throat, then you should ask clarifying questions like:
'When did the sore throat start?'
'Have you had any other symptoms like coughing or congestion?'
'Is there swelling or redness in your throat?'
'Do you feel pain in your throat?'
Make sure to sure to ask about related symptoms. For example, if you have a sore throat ask about a fever, coughing, congestion, head aches, body aches, and general weakness. These are not the only examples of related symptoms use your judgment to ask about possibly related symptoms.

-You should continue this portion for about 1-2 minutes in conversation to gather enough information for a report.

Maintain Clarity and Focus:
Your questions should be straightforward, avoiding medical jargon that might confuse the patient.
Listen attentively to the patient's responses to guide the flow of the conversation naturally toward the most relevant topics.

Adhere to the One-Question Rule:
After each patient response, pause and reflect on the information provided.
Formulate ONE specific follow-up question that naturally extends from the patient's last answer. This disciplined approach ensures that the conversation remains focused and the patient does not feel overwhelmed.

Key Rules and Guidelines:
-One Question at a Time: To prevent overwhelming the patient and to maintain a smooth conversational flow, limit yourself to asking one question per interaction. Await the patient's response before proceeding to the next question.
-Closing the Conversation: Once you've gathered all necessary information, it's important to conclude the conversation gracefully and reassuringly. Use the following exact phrase to end the conversation: "Thank you for your time, we'll see you in the office later today." This statement should not be altered in any way and serves as a clear signal that the assessment is complete.

Tips for success:
Adaptability: Be prepared to adjust the direction of the conversation based on the patient's responses. This might mean revisiting earlier topics or introducing new questions as needed.
Empathy and Patience: Always approach the conversation with empathy and patience, understanding that patients may have concerns or anxieties about their health.

Mandatory Conversation Closure:

Once you've gathered sufficient information, conclude the assessment with the exact phrase: "Thank you for your time, we'll see you in the office later today." This specific sentence signals the end of the conversation and must be used verbatim.

By following these guidelines, you will contribute significantly to improving patient care efficiency and experience. Your role as a Virtual Nurse is pivotal in Breezy's mission to enhance primary care through technology.
"""



class MedicalChatbot:
    """
    A class to manage a virtual nurse-patient conversation, generate responses,
    and create a patient assessment report.
    """

    def __init__(self):
        self.context = [{"role": "system", "content": message_prompt}]
        self.finished = False
        self.patient_info = {
            'Subjective': '',
            'Objective': '',
            'Analysis': '',
            'Plan': '',
            'Implementation': '',
            'Evaluation': ''
        }

        self.initial_questions_answers = ""
    
    def handle_initial_questions(self):
        """
        Collects answers to predefined questions before starting the intelligent probing.
        Returns the answer to the last question to kick-start the LLM-based probing.
        """
        questions = [
            "Are you a new patient?",
            "What is your name?",
            "What is your approximate height?",
            "What is your approximate weight?",
            "Are you currently taking any medications?",
            "Have you had any recent surgeries?",
            "Do you have any known drug allergies?",
            "Finally, what are you in for today?"
        ]
        # Initialize a dictionary to store questions and answers
        self.initial_questions_dict = {}
        last_answer = ""
        for question in questions:
            print(question)
            answer = input("You: ")
            self.initial_questions_dict[question] = answer
            if question.startswith("Finally"):
                last_answer = answer  # Save the last answer for transition

        # The 'Subjective' part will be updated in create_report or extract_and_save_report
        return last_answer
    
    def should_stop(self, message):
        """
        Determine if the conversation should be ended based on the user's message.
        """
        if "we'll see you in the office later today" in message.lower():
            self.finished = True


    def generate_response(self, message):
        """
        Generate a response to the user's message using the OpenAI API.
        """
        self.context.append({'role': 'user', 'content': message})
        response = client.chat.completions.create(
            model="gpt-4",
            messages=self.context
        )
        assistant_message = response.choices[0].message.content
        self.context.append({'role': 'assistant', 'content': assistant_message})
        return assistant_message

    def update_patient_info(self, category, content):
        """
        Update patient info for the report.
        """
        self.patient_info[category] += f"{content}\n"

    def get_full_conversation(self):
        """
        Return the full conversation history.
        """
        return self.context

    def create_report(self):
        chat_history = self.get_full_conversation()[1:] 
        # Build the prompt for the report
        new_prompt = """\nPHASE 3- Documentation Synthesis:
Document the patient's responses and any additional relevant information, as a nurse would for a doctor to use.

*rules*
-You have examples of what a nurses notes look like. Follow the 'SOAP' model for writing notes, you also have access to this in your knowledge.
-You should be very careful filling the O-Objective part of the note as most times you will not have the information to fill this out.
-Keep the notes easy and quick to read for a doctor. Ensure the notes get created in a Microsoft word document.
-You MUST use your knowledge of how to write good nursing notes and nursing note examples to format your notes properly.
-You MUST create a word document of the notes to deliver your final notes.
- You MUST use BULLET POINTS to allow for quick reading by doctor. They should not be repetitive
- For the implementation tab, you MUST not talk about communicating to a healthcare provider because they are visiting a doctor already

Here is an example document:

Subjective: 
- Patient reports a dull ache in the left shoulder, which started about a month ago
- The pain intensifies during workouts, especially during bench press. No previous injuries reported.
- No swelling, redness, or warmth in the area. Patient mentions a clicking sound when raising the arm laterally.
Objective: 
- Patient has been using ice and ibuprofen for pain management. 
- Ibuprofen provides temporary relief, and is needed approximately twice a day.
Analysis: 
- The symptoms suggest a potential joint issue in the left shoulder, possibly related to exercise. 
- The clicking sound during movement may indicate a ligament or tendon problem.
Plan: 

- Possible imaging tests or referral to physical therapy may be considered based on the assessment.
Implementation:
- Advised the patient on the importance of consulting a healthcare provider
- Discussed the potential need for further diagnostic tests or physical therapy.
Evaluation: 
- To be determined based on patient's follow-up with healthcare provider and any subsequent treatment.

Make sure to include the implementation and evaluation.

Here is the chat history to base this off of below: \n"""

        
        chat_history_string = ""
        for message in chat_history:
            role = message['role'].capitalize()
            content = message['content'].replace('\n', ' ')  # Replace newlines with spaces
            chat_history_string += f"{role}: {content}\n"
        
        new_prompt += chat_history_string
        
        try:
            # Call the OpenAI API to generate the report
            response =  client.chat.completions.create(
                model="gpt-4",
                messages=[{"role" : "system", "content" : new_prompt}]
            )
            return response
        except Exception as e:
            print(f"An error occurred {str(e)}")
            return ""

    def extract_and_save_report(self, report_content):
        """
        Extract the report content from the response and save it as a Word document,
        including initial questions and answers at the top.
        """
        try:
            doc = Document()
            doc.add_heading('Patient Assessment Report', 0)

            # Add initial questions and answers
            for question, answer in self.initial_questions_dict.items():
                doc.add_paragraph(f"{question}: {answer}")

            # Add a separator before the rest of the report
            doc.add_paragraph("\n---\n")

            # Add the report content
            doc.add_paragraph(report_content)
            file_path = 'Patient_Assessment_Report.docx'  # Adjust path as necessary
            doc.save(file_path)
            return file_path
        except Exception as e:
            return f"An error occurred: {str(e)}"

def main():
    print("Hello, I'm your virtual nurse assistant. Let's start with some basic questions.")
    bot = MedicalChatbot()
    last_initial_answer = bot.handle_initial_questions()  # Collect answers and get the last answer

    # Directly use the last initial answer to start the LLM-based probing
    if last_initial_answer:
        response = bot.generate_response(last_initial_answer)
        print("Virtual Nurse:", response)

    while not bot.finished:
        user_input = input("You: ")
        response = bot.generate_response(user_input)
        print("Virtual Nurse:", response)
        bot.should_stop(response)


    if bot.finished:
        report_content = bot.create_report().choices[0].message.content  # Assuming create_report returns a response object
        file_path = bot.extract_and_save_report(report_content)
        print(f"Report saved to: {file_path}")

if __name__ == "__main__":
    main()